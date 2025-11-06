from flask import Flask, request, send_file, abort, jsonify
from io import BytesIO
from docx import Document
from docx.shared import Pt
from datetime import datetime
from decimal import Decimal, ROUND_HALF_UP
from flask_cors import CORS
import os, re

API_KEY = os.getenv('PIPE_DEMAND_API_KEY', 'YOUR_SECRET_KEY')

app = Flask(__name__)
# Allow calls from your Apps Script web app (and localhost for testing)
CORS(app, resources={r"/demand-letter": {"origins": ["*"]}})  # tighten later to your domain if desired

LETTER_HEAD = ['Servicing and Collections', 'Pipe Advance LLC']

TEMPLATE = {
  'header': (
    '{Business Name}\n{Business Address}\nUnited States of America\n\n'
    'SENT VIA EMAIL ON TODAY’S DATE {Today}\n\n'
    'Re: Demand for Payment - Pipe Merchant Cash Advance\n\n'
    'Dear {ContactName},\n\n'
  ),
  'body': (
    'This is our last attempt and FINAL WARNING to seek payment for {BusinessName}’s merchant cash advance ("MCA") '
    'before we seek all legal remedies available to us. {BusinessName} ("you") entered into an MCA Agreement '
    '("Agreement") with Pipe Advancel LLC (the "Company") dated {EffectiveDate} (the "Effective Date") for an MCA '
    'in the total amount of {TotalAdvancePlusFeeFmt} (consisting of a MCA advance of {AdvanceAmountFmt} and a fee of {FeeFmt}).\n\n'
    'Since {DefaultDate}, {BusinessName} has failed to comply with its terms, by generating revenue and failing to deliver '
    'and/or preventing Pipe from receiving its share of revenue payments. As of {Today}, {BusinessName} has had '
    '{TotalRevenueFmt} in revenue payments of which {RRPercentFmt} ({RRAmountFmt}) are payable to Pipe under the terms of the '
    'Agreement. We have only received {SuccessfulPaymentsFmt} towards your Total Advance Amount. The last payment to Pipe was on '
    '{LastPaymentDate}.\n\n'
    'Your failure to pay Pipe the agreed upon percentage of revenue {PercentOrAmountDueFmt}, is a breach of the Agreement. '
    'We have attempted to contact you and resolve this issue informally multiple times. Despite Pipe’s continuous efforts to '
    'resolve this issue, we have not received a payment.\n\n'
    'If a payment of {ShortfallFmt} is not received within 3 business days of receipt of this letter, we will seek all remedies '
    'available to us under the Agreement, including referring this matter to a third-party collections firm or seeking appropriate '
    'legal action. You may also be held liable and subject to additional fees incurred by Pipe in an attempt to pursue these payments.\n\n'
    'We urge you to treat this matter with the utmost urgency and to cooperate fully in resolving this breach amicably.\n\n'
  ),
  'footer': (
    'Please contact our Servicing and Collections Manager, William, at william@pipe.com immediately within the next 3 business days.\n\n'
    'Thank you for your immediate attention to this critical issue.\n\n'
    'Servicing and Collections\nPipe Advance LLC\n'
  )
}

CURRENCY_RE = re.compile(r'[^0-9.-]')

def money_to_decimal(s, field):
  if s is None or str(s).strip() == '':
    return None
  try:
    from decimal import Decimal
    val = Decimal(CURRENCY_RE.sub('', str(s))).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
    if val < Decimal('0'):
      raise ValueError
    return val
  except Exception:
    raise ValueError(f"Invalid currency for {field}: '{s}'")

def percent_to_decimal(s, field):
  if s is None or str(s).strip() == '':
    return None
  try:
    from decimal import Decimal
    s = str(s).strip().replace('%','')
    val = Decimal(s)
    if val < 0 or val > 100:
      raise ValueError
    return val
  except Exception:
    raise ValueError(f"Invalid percentage for {field}: '{s}' (expected 0–100)")

def parse_date(s, field):
  if not s:
    raise ValueError(f"Missing required date: {field}")
  try:
    dt = datetime.strptime(s.strip(), '%b %d %Y')
    return dt.strftime('%b %d %Y')
  except Exception:
    raise ValueError(f"Invalid date format for {field}: '{s}' (use MMM DD YYYY, e.g., Nov 06 2025)")

def fmt_money(d): return f"${d:,.2f}"
def require(value, field):
  if value is None or str(value).strip() == '':
    raise ValueError(f"Missing required field: {field}")
  return value

@app.get("/healthz")
def healthz():
  return {"status": "ok"}

@app.post('/demand-letter')
def demand_letter():
  if request.headers.get('X-API-KEY') != API_KEY:
    return abort(401)

  data = request.get_json(force=True) or {}
  try:
    business_name = require(data.get('business_name'), 'Business Name')
    business_address = require(data.get('business_address'), 'Business Address')
    contact_name = data.get('contact_name') or 'Client'

    effective_date = parse_date(require(data.get('effective_date'), 'Effective Date'), 'Effective Date')
    today = parse_date(data.get('today') or datetime.utcnow().strftime('%b %d %Y'), 'Today')
    default_date = parse_date(require(data.get('default_date'), 'Date of Default Event'), 'Date of Default Event')
    last_payment_date = parse_date(require(data.get('last_payment_date'), 'Date of Last Payment'), 'Date of Last Payment')

    total_advance_plus_fee = money_to_decimal(require(data.get('total_advance_plus_fee'), 'Total Advance + Fee'), 'Total Advance + Fee')
    advance_amount = money_to_decimal(require(data.get('advance_amount'), 'Advance Amount'), 'Advance Amount')
    fee = money_to_decimal(require(data.get('fee'), 'Fee'), 'Fee')
    total_revenue = money_to_decimal(require(data.get('total_revenue'), 'Total Revenue Since Agreement to Today'), 'Total Revenue Since Agreement to Today')
    rr_percent = percent_to_decimal(require(data.get('rr_percent'), 'Revenue Share Percentage (RR%)'), 'Revenue Share Percentage (RR%)')

    rr_amount = money_to_decimal(data.get('rr_amount'), 'Calculated % of Revenue Payable to Pipe ($)')
    successful_payments = money_to_decimal(require(data.get('successful_payments'), 'Amount of Successful Payments ($)'), 'Amount of Successful Payments ($)')
    percent_or_amount_due = money_to_decimal(require(data.get('percent_or_amount_due'), 'Payment Percentage or Amount Due ($% of Revenue Amount)'), 'Payment Percentage or Amount Due ($% of Revenue Amount)')
    shortfall = money_to_decimal(data.get('shortfall'), 'Shortfall (Amount Needed to Cure Default) ($)')

    # Auto-calcs
    calc_rr_amount = (total_revenue * rr_percent / Decimal('100')).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
    if rr_amount is None or rr_amount != calc_rr_amount:
      rr_amount = calc_rr_amount

    if shortfall is None:
      shortfall = (rr_amount - successful_payments).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
      if shortfall < 0:
        shortfall = Decimal('0.00')

    m = {
      'BusinessName': business_name,
      'Business Address': business_address,
      'ContactName': contact_name,
      'EffectiveDate': effective_date,
      'Today': today,
      'DefaultDate': default_date,
      'LastPaymentDate': last_payment_date,
      'TotalAdvancePlusFeeFmt': fmt_money(total_advance_plus_fee),
      'AdvanceAmountFmt': fmt_money(advance_amount),
      'FeeFmt': fmt_money(fee),
      'TotalRevenueFmt': fmt_money(total_revenue),
      'RRPercentFmt': f"{rr_percent}%",
      'RRAmountFmt': fmt_money(rr_amount),
      'SuccessfulPaymentsFmt': fmt_money(successful_payments),
      'PercentOrAmountDueFmt': fmt_money(percent_or_amount_due),
      'ShortfallFmt': fmt_money(shortfall),
    }

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    doc.add_paragraph(TEMPLATE['header'].format(**{
      'Business Name': business_name,
      'Business Address': business_address,
      'Today': today,
      'ContactName': contact_name
    }))
    doc.add_paragraph(TEMPLATE['body'].format(**m))
    doc.add_paragraph(TEMPLATE['footer'])
    for line in LETTER_HEAD:
      doc.add_paragraph(line)

    buf = BytesIO()
    doc.save(buf); buf.seek(0)
    safe_name = re.sub(r'[\\s]+', '_', business_name)
    return send_file(buf,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                     as_attachment=True,
                     download_name=f"Demand_Letter_{safe_name}.docx")
  except ValueError as ve:
    return jsonify({'error': str(ve)}), 400
